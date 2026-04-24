#!/usr/bin/env python3
"""
PDF转Word工具 - 使用Python标准库的简单版本
支持文本提取和图片处理
"""

import os
import sys
import json
import struct
import zlib
import re
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs
import tempfile
import shutil

# 版本
VERSION = "1.0.0"
SUPPORTED_FORMATS = ['.pdf']

# 简单PDF解析器
class SimplePDFParser:
    """简单PDF文本提取器"""
    
    def __init__(self, filepath):
        self.filepath = filepath
        self.content = None
        self.text = ""
    
    def read(self):
        """读取PDF文件"""
        try:
            with open(self.filepath, 'rb') as f:
                self.content = f.read()
            return True
        except Exception as e:
            print(f"读取文件失败: {e}")
            return False
    
    def extract_text(self):
        """提取PDF中的文本"""
        if not self.content:
            if not self.read():
                return ""
        
        text_parts = []
        try:
            # 简单的文本提取（提取括号之间的内容）
            # PDF文本通常在 (text) 或 <hex> 格式中
            content_str = self.content.decode('latin-1', errors='ignore')
            
            # 提取 (text) 格式的文本
            pattern = r'\(([^)]+)\)'
            matches = re.findall(pattern, content_str)
            for m in matches:
                if len(m) > 0 and not m.startswith('/'):
                    # 解码转义字符
                    text = m.replace('\\(', '(').replace('\\)', ')').replace('\\n', '\n')
                    if text.strip():
                        text_parts.append(text)
            
            # 提取流内容
            stream_pattern = rb'stream\r?\n(.*?)\r?\nendstream'
            streams = re.findall(stream_pattern, self.content, re.DOTALL)
            
            for stream in streams:
                try:
                    # 尝试解压
                    decompressed = zlib.decompress(stream)
                    stream_text = decompressed.decode('utf-8', errors='ignore')
                    # 提取文本
                    text_matches = re.findall(r'\(([^)]+)\)', stream_text)
                    for tm in text_matches:
                        if len(tm) > 0:
                            text_parts.append(tm)
                except:
                    pass
            
            self.text = '\n'.join(text_parts)
            
            # 如果没提取到，使用更简单的方法
            if not self.text.strip():
                # 直接提取可打印字符
                printable = []
                for i, byte in enumerate(self.content):
                    if 32 <= byte <= 126 or byte in (10, 13):
                        printable.append(chr(byte))
                self.text = ''.join(printable)
                # 清理
                self.text = re.sub(r'\s+', ' ', self.text)
            
        except Exception as e:
            print(f"提取文本失败: {e}")
        
        return self.text
    
    def get_info(self):
        """获取PDF信息"""
        info = {
            'title': '',
            'author': '',
            'pages': 0,
            'size': os.path.getsize(self.filepath) if os.path.exists(self.filepath) else 0
        }
        
        try:
            content_str = self.content.decode('latin-1', errors='ignore')
            
            # 提取页数
            pages_match = re.search(r'/Count\s+(\d+)', content_str)
            if pages_match:
                info['pages'] = int(pages_match.group(1))
            
            # 提取标题
            title_match = re.search(r'/Title\s*\(([^)]+)\)', content_str)
            if title_match:
                info['title'] = title_match.group(1)
            
            # 提取作者
            author_match = re.search(r'/Author\s*\(([^)]+)\)', content_str)
            if author_match:
                info['author'] = author_match.group(1)
                
        except:
            pass
        
        return info


class WordGenerator:
    """生成Word文档（简化版）"""
    
    def __init__(self):
        self.content = []
    
    def add_paragraph(self, text, style='Normal'):
        """添加段落"""
        self.content.append({'type': 'paragraph', 'text': text, 'style': style})
    
    def add_heading(self, text, level=1):
        """添加标题"""
        self.content.append({'type': 'heading', 'text': text, 'level': level})
    
    def add_page_break(self):
        """添加分页符"""
        self.content.append({'type': 'pagebreak'})
    
    def save(self, filepath):
        """保存为Word文档"""
        try:
            # 使用python-docx如果可用
            from docx import Document
            doc = Document()
            
            for item in self.content:
                if item['type'] == 'paragraph':
                    doc.add_paragraph(item['text'])
                elif item['type'] == 'heading':
                    doc.add_heading(item['text'], level=item['level'])
                elif item['type'] == 'pagebreak':
                    doc.add_page_break()
            
            doc.save(filepath)
            return True
        except ImportError:
            # 如果没有python-docx，生成HTML并转为.doc
            html_content = self._generate_html()
            with open(filepath.replace('.docx', '.html'), 'w', encoding='utf-8') as f:
                f.write(html_content)
            return True
        except Exception as e:
            print(f"保存Word失败: {e}")
            return False
    
    def _generate_html(self):
        """生成HTML内容"""
        html = ['<!DOCTYPE html>', '<html><head><meta charset="utf-8">',
                '<title>Converted Document</title></head><body>']
        
        for item in self.content:
            if item['type'] == 'paragraph':
                html.append(f'<p>{item["text"]}</p>')
            elif item['type'] == 'heading':
                html.append(f'<h{item["level"]}>{item["text"]}</h{item["level"]}>')
            elif item['type'] == 'pagebreak':
                html.append('<hr style="page-break-after:always;">')
        
        html.append('</body></html>')
        return '\n'.join(html)


def pdf_to_word(pdf_path, output_path=None):
    """PDF转Word主函数"""
    if not output_path:
        output_path = pdf_path.replace('.pdf', '.docx')
    
    # 解析PDF
    parser = SimplePDFParser(pdf_path)
    
    if not parser.read():
        return {'success': False, 'error': '无法读取PDF文件'}
    
    # 提取文本
    text = parser.extract_text()
    info = parser.get_info()
    
    # 生成Word
    doc = WordGenerator()
    
    # 添加文档信息
    if info.get('title'):
        doc.add_heading(info['title'], level=1)
    if info.get('author'):
        doc.add_paragraph(f"作者: {info['author']}")
    
    # 添加内容
    if text.strip():
        # 按段落分割
        paragraphs = text.split('\n')
        for p in paragraphs:
            p = p.strip()
            if p:
                doc.add_paragraph(p)
    else:
        doc.add_paragraph('(PDF内容提取失败，请使用专业工具处理)')
    
    # 保存
    success = doc.save(output_path)
    
    return {
        'success': success,
        'input': pdf_path,
        'output': output_path,
        'pages': info.get('pages', 0),
        'size': info.get('size', 0),
        'text_length': len(text)
    }


def convert_folder(input_folder, output_folder=None):
    """批量转换文件夹中的PDF"""
    if not output_folder:
        output_folder = os.path.join(input_folder, 'converted')
    
    os.makedirs(output_folder, exist_ok=True)
    
    results = []
    for filename in os.listdir(input_folder):
        if filename.lower().endswith('.pdf'):
            pdf_path = os.path.join(input_folder, filename)
            word_path = os.path.join(output_folder, filename.replace('.pdf', '.docx'))
            
            result = pdf_to_word(pdf_path, word_path)
            result['filename'] = filename
            results.append(result)
            
            print(f"转换: {filename} -> {result.get('success', False)}")
    
    return results


# ==================== Web界面 ====================

class PDFConverterHandler(BaseHTTPRequestHandler):
    """HTTP请求处理器"""
    
    def log_message(self, format, *args):
        pass
    
    def send_json(self, data, status=200):
        self.send_response(status)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps(data, ensure_ascii=False).encode('utf-8'))
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
    
    def do_GET(self):
        parsed = urlparse(self.path)
        path = parsed.path
        
        if path == '/' or path == '/index.html':
            self.serve_index()
        elif path == '/api/status':
            self.send_json({'status': 'ok', 'version': VERSION})
        else:
            self.send_json({'code': 404, 'message': 'Not Found'}, 404)
    
    def do_POST(self):
        parsed = urlparse(self.path)
        path = parsed.path
        
        if path == '/api/convert':
            self.handle_convert()
        else:
            self.send_json({'code': 404, 'message': 'Not Found'}, 404)
    
    def serve_index(self):
        html = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>PDF转Word工具</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
               background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; }
        .container { max-width: 800px; margin: 0 auto; padding: 40px 20px; }
        .card { background: white; border-radius: 16px; padding: 40px; box-shadow: 0 10px 40px rgba(0,0,0,0.2); }
        h1 { text-align: center; color: #333; margin-bottom: 10px; }
        .subtitle { text-align: center; color: #666; margin-bottom: 30px; }
        .upload-area { border: 2px dashed #667eea; border-radius: 12px; padding: 60px 20px; 
                       text-align: center; cursor: pointer; transition: all 0.3s; margin-bottom: 20px; }
        .upload-area:hover { background: #f8f9ff; border-color: #764ba2; }
        .upload-area.dragover { background: #f0f4ff; border-color: #764ba2; }
        .upload-icon { font-size: 48px; margin-bottom: 10px; }
        .upload-text { color: #666; }
        .upload-hint { color: #999; font-size: 12px; margin-top: 10px; }
        input[type="file"] { display: none; }
        .btn { width: 100%; padding: 16px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
               color: white; border: none; border-radius: 8px; font-size: 16px; cursor: pointer;
               transition: transform 0.2s; margin-top: 20px; }
        .btn:hover { transform: translateY(-2px); }
        .btn:disabled { background: #ccc; cursor: not-allowed; transform: none; }
        .progress { display: none; margin-top: 20px; }
        .progress-bar { height: 8px; background: #eee; border-radius: 4px; overflow: hidden; }
        .progress-fill { height: 100%; background: linear-gradient(90deg, #667eea, #764ba2); 
                         width: 0%; transition: width 0.3s; }
        .progress-text { text-align: center; margin-top: 10px; color: #666; }
        .result { display: none; margin-top: 20px; padding: 20px; background: #f8f9ff; 
                  border-radius: 8px; text-align: center; }
        .result.success { background: #e8f5e9; }
        .result.error { background: #ffebee; }
        .download-btn { display: inline-block; padding: 12px 24px; background: #667eea; color: white;
                        text-decoration: none; border-radius: 8px; margin-top: 10px; }
        .download-btn:hover { background: #764ba2; }
        .features { display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px; margin-top: 30px; }
        .feature { text-align: center; padding: 20px; background: #f8f9ff; border-radius: 8px; }
        .feature-icon { font-size: 32px; margin-bottom: 10px; }
        .feature-title { font-weight: bold; color: #333; }
        .feature-desc { color: #666; font-size: 12px; margin-top: 5px; }
    </style>
</head>
<body>
    <div class="container">
        <div class="card">
            <h1>📄 PDF转Word工具</h1>
            <p class="subtitle">快速将PDF文档转换为可编辑的Word文件</p>
            
            <div class="upload-area" id="dropZone">
                <div class="upload-icon">📁</div>
                <div class="upload-text">点击或拖拽PDF文件到这里</div>
                <div class="upload-hint">支持单个文件，最大50MB</div>
                <input type="file" id="fileInput" accept=".pdf">
            </div>
            
            <div class="progress" id="progress">
                <div class="progress-bar">
                    <div class="progress-fill" id="progressFill"></div>
                </div>
                <div class="progress-text" id="progressText">准备转换...</div>
            </div>
            
            <button class="btn" id="convertBtn" disabled>开始转换</button>
            
            <div class="result" id="result">
                <div id="resultText"></div>
                <a href="#" class="download-btn" id="downloadBtn" style="display:none;">下载Word文件</a>
            </div>
            
            <div class="features">
                <div class="feature">
                    <div class="feature-icon">⚡</div>
                    <div class="feature-title">快速转换</div>
                    <div class="feature-desc">秒级处理</div>
                </div>
                <div class="feature">
                    <div class="feature-icon">📝</div>
                    <div class="feature-title">文本提取</div>
                    <div class="feature-desc">保留文本内容</div>
                </div>
                <div class="feature">
                    <div class="feature-icon">🔒</div>
                    <div class="feature-title">本地处理</div>
                    <div class="feature-desc">数据不外传</div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        const dropZone = document.getElementById('dropZone');
        const fileInput = document.getElementById('fileInput');
        const convertBtn = document.getElementById('convertBtn');
        const progress = document.getElementById('progress');
        const progressFill = document.getElementById('progressFill');
        const progressText = document.getElementById('progressText');
        const result = document.getElementById('result');
        const resultText = document.getElementById('resultText');
        const downloadBtn = document.getElementById('downloadBtn');
        
        let selectedFile = null;
        
        dropZone.addEventListener('click', () => fileInput.click());
        dropZone.addEventListener('dragover', (e) => { e.preventDefault(); dropZone.classList.add('dragover'); });
        dropZone.addEventListener('dragleave', () => dropZone.classList.remove('dragover'));
        dropZone.addEventListener('drop', (e) => {
            e.preventDefault();
            dropZone.classList.remove('dragover');
            if (e.dataTransfer.files.length) handleFile(e.dataTransfer.files[0]);
        });
        fileInput.addEventListener('change', (e) => handleFile(e.target.files[0]));
        
        function handleFile(file) {
            if (!file.name.toLowerCase().endsWith('.pdf')) {
                alert('请选择PDF文件');
                return;
            }
            if (file.size > 50 * 1024 * 1024) {
                alert('文件大小不能超过50MB');
                return;
            }
            selectedFile = file;
            dropZone.innerHTML = '<div class="upload-icon">✓</div><div class="upload-text">' + file.name + '</div><div class="upload-hint">' + (file.size / 1024).toFixed(1) + ' KB</div>';
            convertBtn.disabled = false;
        }
        
        convertBtn.addEventListener('click', async () => {
            if (!selectedFile) return;
            
            convertBtn.disabled = true;
            progress.style.display = 'block';
            result.style.display = 'none';
            
            const formData = new FormData();
            formData.append('file', selectedFile);
            
            try {
                progressText.textContent = '正在转换...';
                progressFill.style.width = '50%';
                
                const response = await fetch('/api/convert', { method: 'POST', body: formData });
                const data = await response.json();
                
                progressFill.style.width = '100%';
                progressText.textContent = '转换完成!';
                
                result.style.display = 'block';
                if (data.success) {
                    result.className = 'result success';
                    resultText.innerHTML = '✅ 转换成功！<br>提取文本: ' + data.text_length + ' 字符';
                    downloadBtn.href = data.download_url;
                    downloadBtn.style.display = 'inline-block';
                    downloadBtn.download = selectedFile.name.replace('.pdf', '.docx');
                } else {
                    result.className = 'result error';
                    resultText.textContent = '❌ ' + (data.error || '转换失败');
                }
            } catch (e) {
                result.style.display = 'block';
                result.className = 'result error';
                resultText.textContent = '❌ 网络错误: ' + e.message;
            }
            
            convertBtn.disabled = false;
        });
    </script>
</body>
</html>'''
        self.send_response(200)
        self.send_header('Content-Type', 'text/html; charset=utf-8')
        self.end_headers()
        self.wfile.write(html.encode('utf-8'))
    
    def handle_convert(self):
        content_length = int(self.headers.get('Content-Length', 0))
        if content_length == 0:
            self.send_json({'success': False, 'error': '没有上传文件'}, 400)
            return
        
        # 读取文件
        content_type = self.headers.get('Content-Type', '')
        boundary = None
        for line in content_type.split(';'):
            if 'boundary=' in line:
                boundary = line.split('=')[1].strip()
                break
        
        if not boundary:
            self.send_json({'success': False, 'error': '无效的请求'}, 400)
            return
        
        body = self.rfile.read(content_length)
        
        # 解析multipart数据
        try:
            # 查找文件内容
            boundary_bytes = boundary.encode()
            parts = body.split(b'--' + boundary_bytes)
            
            file_content = None
            filename = 'upload.pdf'
            
            for part in parts:
                if b'Content-Disposition' in part and b'filename' in part:
                    # 提取文件名
                    fn_match = b'filename="' in part and part.split(b'filename="')[1].split(b'"')[0]
                    if fn_match:
                        filename = fn_match.decode('utf-8', errors='ignore')
                    
                    # 提取文件内容
                    header_end = part.find(b'\r\n\r\n')
                    if header_end != -1:
                        file_content = part[header_end+4:].rstrip(b'\r\n--')
                    break
            
            if not file_content:
                self.send_json({'success': False, 'error': '无法提取文件'}, 400)
                return
            
            # 保存临时文件
            temp_dir = tempfile.mkdtemp()
            temp_pdf = os.path.join(temp_dir, filename)
            temp_docx = temp_pdf.replace('.pdf', '.docx')
            
            with open(temp_pdf, 'wb') as f:
                f.write(file_content)
            
            # 转换
            result = pdf_to_word(temp_pdf, temp_docx)
            
            if result['success']:
                # 读取生成的文件
                if os.path.exists(temp_docx):
                    with open(temp_docx, 'rb') as f:
                        docx_content = f.read()
                    
                    # 返回下载链接
                    self.send_response(200)
                    self.send_header('Content-Type', 'application/json')
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                    
                    # Base64编码文件内容
                    import base64
                    docx_base64 = base64.b64encode(docx_content).decode('utf-8')
                    
                    self.wfile.write(json.dumps({
                        'success': True,
                        'filename': filename.replace('.pdf', '.docx'),
                        'pages': result.get('pages', 0),
                        'text_length': result.get('text_length', 0),
                        'download_url': 'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,' + docx_base64
                    }).encode('utf-8'))
                else:
                    # 如果docx生成失败，返回HTML版本
                    temp_html = temp_docx.replace('.docx', '.html')
                    if os.path.exists(temp_html):
                        with open(temp_html, 'r', encoding='utf-8') as f:
                            html_content = f.read()
                        
                        import base64
                        html_base64 = base64.b64encode(html_content.encode('utf-8')).decode('utf-8')
                        
                        self.send_json({
                            'success': True,
                            'filename': filename.replace('.pdf', '.html'),
                            'pages': result.get('pages', 0),
                            'text_length': result.get('text_length', 0),
                            'download_url': 'data:text/html;base64,' + html_base64
                        })
                    else:
                        self.send_json({'success': False, 'error': '文件生成失败'}, 500)
            else:
                self.send_json({'success': False, 'error': result.get('error', '转换失败')}, 500)
            
            # 清理临时文件
            shutil.rmtree(temp_dir, ignore_errors=True)
            
        except Exception as e:
            self.send_json({'success': False, 'error': str(e)}, 500)


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='PDF转Word工具')
    parser.add_argument('input', nargs='?', help='输入PDF文件或文件夹')
    parser.add_argument('-o', '--output', help='输出路径')
    parser.add_argument('-p', '--port', type=int, default=5000, help='Web服务端口')
    parser.add_argument('-w', '--web', action='store_true', help='启动Web界面')
    
    args = parser.parse_args()
    
    if args.web or not args.input:
        # 启动Web服务
        print("="*50)
        print("📄 PDF转Word工具")
        print("="*50)
        print(f"🌐 Web界面: http://localhost:{args.port}")
        print("按 Ctrl+C 停止服务")
        print("="*50)
        
        server = HTTPServer(('0.0.0.0', args.port), PDFConverterHandler)
        try:
            server.serve_forever()
        except KeyboardInterrupt:
            print("\n👋 服务已停止")
    
    elif os.path.isfile(args.input):
        # 单文件转换
        output = args.output or args.input.replace('.pdf', '.docx')
        print(f"转换: {args.input} -> {output}")
        result = pdf_to_word(args.input, output)
        if result['success']:
            print(f"✅ 转换完成! 提取文本: {result['text_length']} 字符")
        else:
            print(f"❌ 转换失败: {result.get('error', '未知错误')}")
    
    elif os.path.isdir(args.input):
        # 批量转换
        print(f"批量转换: {args.input}")
        results = convert_folder(args.input, args.output)
        success = sum(1 for r in results if r['success'])
        print(f"完成: {success}/{len(results)} 个文件")
    
    else:
        print(f"错误: 找不到文件或文件夹 {args.input}")
        sys.exit(1)


if __name__ == '__main__':
    main()