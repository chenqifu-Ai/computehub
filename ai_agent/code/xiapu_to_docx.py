#!/usr/bin/env python3
"""OCR 结果 → Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)

# ── 标题 ──
title = doc.add_heading('霞浦县美丽海湾建设与产业一体化开发（EOD）项目', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('股 东 协 议')
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()  # spacing

with open('/tmp/xiapu_ocr/霞浦股东协议_OCR结果.txt') as f:
    text = f.read()

# ── Split pages ──
pages = re.split(r'=+\n+第\s*(\d+)\s*页\s*=+', text)
# pages[0] is page 1's content (before first separator)
# pages[1] = page number, pages[2] = content, alternate

# First content block
if pages:
    lines = pages[0].strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)

# Remaining pages
for i in range(1, len(pages) - 1, 2):
    page_num = pages[i].strip()
    content = pages[i+1].strip()
    
    # Page header
    doc.add_paragraph('')
    h = doc.add_heading(f'第 {page_num} 页', level=2)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Detect directory/chapter headings
        if line.startswith('目 录') or re.match(r'^第[一二三四五六七八九十]+章', line):
            p = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

out_path = '/tmp/xiapu_ocr/霞浦股东协议.docx'
doc.save(out_path)
print(f'✅ Word 文档已生成: {out_path}')
print(f'   大小: {__import__("os").path.getsize(out_path)/1024:.0f}KB')